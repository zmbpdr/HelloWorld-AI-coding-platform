"""文件导入解析服务 — Word/PDF → Markdown

使用 python-docx 解析 Word 文档，PyMuPDF 解析 PDF 文档，
将内容转换为 Markdown 格式，提取内嵌图片并上传。
"""

import io
import os
import uuid
from datetime import datetime
from pathlib import Path

from app.config import settings


def _save_inline_image(image_bytes: bytes, ext: str = "png") -> str:
    """将解析出的内嵌图片保存到上传目录，返回可访问 URL"""
    date_dir = datetime.now().strftime("%Y%m%d")
    filename = f"{uuid.uuid4().hex}.{ext}"
    relative_dir = Path(date_dir)
    upload_dir = Path(settings.UPLOAD_DIR) / relative_dir
    upload_dir.mkdir(parents=True, exist_ok=True)
    file_path = upload_dir / filename
    file_path.write_bytes(image_bytes)
    return f"/uploads/{relative_dir.as_posix()}/{filename}"


def parse_word_to_markdown(file_bytes: bytes) -> dict:
    """解析 Word 文档内容为 Markdown

    Args:
        file_bytes: Word 文件的原始字节

    Returns:
        {"markdown": str, "images": int} — Markdown 内容和提取的图片数量
    """
    try:
        from docx import Document
    except ImportError:
        return {"markdown": "", "error": "python-docx 未安装，请运行 pip install python-docx"}

    doc = Document(io.BytesIO(file_bytes))
    lines: list[str] = []
    image_count = 0

    for para in doc.paragraphs:
        if not para.text.strip() and not para.runs:
            lines.append("")
            continue

        # 处理段落样式
        style_name = para.style.name if para.style else ""

        # 构建段落文本（处理内联格式和图片）
        para_parts: list[str] = []
        for run in para.runs:
            text = run.text
            if not text:
                # 检查是否有内嵌图片
                if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    # 图片通过 drawing 元素嵌入（较复杂，提取 blip 引用）
                    blips = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for blip in blips:
                        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id:
                            try:
                                img_part = doc.part.related_parts[embed_id]
                                img_bytes = img_part.blob
                                img_ext = img_part.content_type.split('/')[-1] if img_part.content_type else 'png'
                                if img_ext == 'jpeg':
                                    img_ext = 'jpg'
                                url = _save_inline_image(img_bytes, img_ext)
                                para_parts.append(f"\n![图片]({url})\n")
                                image_count += 1
                            except Exception:
                                para_parts.append("\n*[图片提取失败]*\n")
                continue

            # 处理内联格式
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"

            para_parts.append(text)

        full_text = "".join(para_parts)

        if not full_text.strip():
            lines.append("")
            continue

        # 根据样式判断标题级别
        if "Heading 1" in style_name or "heading 1" in style_name:
            lines.append(f"# {full_text}")
        elif "Heading 2" in style_name or "heading 2" in style_name:
            lines.append(f"## {full_text}")
        elif "Heading 3" in style_name or "heading 3" in style_name:
            lines.append(f"### {full_text}")
        elif "Heading 4" in style_name or "heading 4" in style_name:
            lines.append(f"#### {full_text}")
        elif "List" in style_name:
            lines.append(f"- {full_text}")
        else:
            lines.append(full_text)

    # 处理表格
    for table in doc.tables:
        lines.append("")
        table_lines: list[str] = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " ") for cell in row.cells]
            table_lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                table_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.extend(table_lines)
        lines.append("")

    markdown = "\n".join(lines)
    return {"markdown": markdown, "images": image_count}


def parse_pdf_to_markdown(file_bytes: bytes) -> dict:
    """解析 PDF 文档内容为 Markdown

    Args:
        file_bytes: PDF 文件的原始字节

    Returns:
        {"markdown": str, "images": int, "pages": int} — Markdown 内容、提取的图片数量和页数
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {"markdown": "", "error": "PyMuPDF 未安装，请运行 pip install PyMuPDF"}

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    lines: list[str] = []
    image_count = 0

    for page_num in range(len(doc)):
        page = doc[page_num]
        lines.append(f"\n### 第 {page_num + 1} 页\n")

        # 提取文本
        text = page.get_text("text")
        if text.strip():
            lines.append(text)

        # 提取图片
        images = page.get_images(full=True)
        for img_info in images:
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                img_ext = base_image["ext"]
                url = _save_inline_image(img_bytes, img_ext)
                lines.append(f"\n![图片]({url})\n")
                image_count += 1
            except Exception:
                lines.append("\n*[图片提取失败]*\n")

    doc.close()

    markdown = "\n".join(lines)
    return {"markdown": markdown, "images": image_count, "pages": len(doc)}