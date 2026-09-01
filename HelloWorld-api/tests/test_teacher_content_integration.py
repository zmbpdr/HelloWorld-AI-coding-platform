"""教师导入教程到学生 AI 上下文的集成回归测试。"""

from io import BytesIO
from types import SimpleNamespace
from unittest.mock import AsyncMock

import pytest
from docx import Document
from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine
from starlette.requests import Request

import app.models  # noqa: F401 - 注册 Base.metadata 中的全部 ORM 模型
from app.database import Base
from app.models.course import Language
from app.models.lesson import Lesson
from app.routers import ai as ai_router
from app.schemas.ai import ChatRequest
from app.services import rag_service
from app.services.file_import_service import parse_word_to_markdown


@pytest.fixture
def anyio_backend():
    return "asyncio"


@pytest.fixture
async def content_session(tmp_path):
    engine = create_async_engine(f"sqlite+aiosqlite:///{tmp_path / 'teacher-content.db'}")
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)

    session_factory = async_sessionmaker(engine, expire_on_commit=False)
    async with session_factory() as session:
        language = Language(name="Python", slug="python")
        session.add(language)
        await session.commit()
        yield session, language
    await engine.dispose()


def _word_bytes() -> bytes:
    document = Document()
    document.add_heading("变量入门", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("变量").bold = True
    paragraph.add_run("用于保存程序中的数据。")
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


@pytest.mark.anyio
async def test_imported_teacher_content_is_retrieved_and_injected_into_ai_chat(content_session, monkeypatch):
    """Word 导入、课程保存、关键词 RAG 回退与 AI 上下文注入必须形成闭环。"""
    session, language = content_session
    imported = parse_word_to_markdown(_word_bytes())
    assert "# 变量入门" in imported["markdown"]
    assert "**变量**" in imported["markdown"]

    lesson = Lesson(
        language_id=language.id,
        title="Python 变量",
        slug="python-variables-test",
        content=imported["markdown"],
        difficulty="beginner",
    )
    unrelated_lesson = Lesson(
        language_id=language.id,
        title="Python 函数",
        slug="python-functions-test",
        content="# 函数\n函数用于复用代码。",
        difficulty="beginner",
    )
    session.add_all([lesson, unrelated_lesson])
    await session.commit()
    await session.refresh(lesson)

    # 不建立外部向量索引时，学生端必须可靠地走数据库关键词检索。
    monkeypatch.setattr(rag_service, "_semantic_search_lesson_context", AsyncMock(return_value=[]))
    captured_context = {}

    async def fallback_search(db, lesson_id, query, limit):
        return await rag_service.search_lesson_context(db, lesson_id, query, limit)

    async def fake_chat_with_ai(message, context):
        captured_context.update(context)
        return "变量可以保存数据。"

    monkeypatch.setattr(ai_router, "search_lesson_context", fallback_search)
    monkeypatch.setattr(ai_router, "ai_limiter", AsyncMock())
    monkeypatch.setattr(ai_router, "chat_with_ai", fake_chat_with_ai)
    monkeypatch.setattr(ai_router, "consume_ai_quota", lambda _user: None)

    response = await ai_router.chat(
        ChatRequest(message="变量", context={"lesson_id": lesson.id}),
        Request({"type": "http", "method": "POST", "path": "/api/v1/ai/chat", "headers": []}),
        current_user=SimpleNamespace(id=99, membership="pro"),
        db=session,
    )

    assert response.reply == "变量可以保存数据。"
    assert captured_context["rag_results"]
    retrieved = captured_context["rag_results"][0]["content"]
    assert "变量" in retrieved
    assert "函数用于复用代码" not in retrieved
