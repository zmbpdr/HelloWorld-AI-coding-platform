"""Word/PDF 教程导入到 Markdown 的核心回归测试。"""

from io import BytesIO

import fitz
from docx import Document

from app.services.file_import_service import parse_pdf_to_markdown, parse_word_to_markdown


def test_word_import_preserves_heading_inline_formatting_and_table():
    document = Document()
    document.add_heading("Python 入门", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("重要").bold = True
    paragraph.add_run("且")
    paragraph.add_run("易学").italic = True
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "知识点"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "变量"
    table.cell(1, 1).text = "存储数据"

    buffer = BytesIO()
    document.save(buffer)
    result = parse_word_to_markdown(buffer.getvalue())

    assert result["images"] == 0
    assert "# Python 入门" in result["markdown"]
    assert "**重要**且*易学*" in result["markdown"]
    assert "| 知识点 | 说明 |" in result["markdown"]
    assert "| --- | --- |" in result["markdown"]


def test_pdf_import_returns_text_and_page_count_after_document_is_closed():
    document = fitz.open()
    first_page = document.new_page()
    first_page.insert_text((72, 72), "First PDF lesson")
    second_page = document.new_page()
    second_page.insert_text((72, 72), "Second PDF lesson")
    pdf_bytes = document.tobytes()
    document.close()

    result = parse_pdf_to_markdown(pdf_bytes)

    assert result["pages"] == 2
    assert result["images"] == 0
    assert "### 第 1 页" in result["markdown"]
    assert "First PDF lesson" in result["markdown"]
    assert "### 第 2 页" in result["markdown"]
    assert "Second PDF lesson" in result["markdown"]
